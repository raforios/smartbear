'''
    Library to create Microsoft Office files
'''
from io import BytesIO
from docx import Document

def create_word_doc(title, content):
    '''
        Function to create Word docx
    '''
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
